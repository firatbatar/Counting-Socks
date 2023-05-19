def save_data(ws, data: dict, start_row: int, start_column: int):
    from openpyxl import load_workbook, Workbook

    idx = 0
    for key in data.keys():
        # Save the keys to first column of the worksheet starting from the second row
        ws.cell(row=start_row + idx, column=start_column, value=key)

        # Save the values to second column of the worksheet starting from the second row
        ws.cell(row=start_row + idx, column=start_column + 1, value=data[key])

        idx += 1


def save_ages(ws, data: dict, row: int, column: int, title_num: str):
    from openpyxl.worksheet.table import Table, TableStyleInfo

    # Data column titles
    ws.cell(row=row, column=column, value="Sock")
    ws.cell(row=row, column=column + 1, value="Age#" + title_num)

    # Save the data
    save_data(ws, data, row + 1, column)

    # Create the table
    tab = Table(displayName="PairsTable" + title_num, ref=f"A{row}"
                                                          f":B{row + len(data)}")
    # Add a default style with striped rows and banded columns
    style = TableStyleInfo(name="TableStyleMedium9", showFirstColumn=False,
                           showLastColumn=False, showRowStripes=True, showColumnStripes=True)
    tab.tableStyleInfo = style  # Apply the style
    # Add table to the worksheet
    ws.add_table(tab)


def determine_bin_bounds(data: list):
    bin_count = 6
    data_min, data_max = min(data) - 0.5, max(data) + 0.5
    step = (data_max - data_min) / bin_count
    if step < 1:
        step = 1
    else:
        dec = step - int(step)
        step = int(step)
        if dec <= 0.5:
            step += 0.5
        else:
            step += 1

    bins = [data_min + step * i for i in range(bin_count + 1)]

    return bins


def plot_histogram(observed: list, title: str = "Title", path: str = "graphs",
                   range_min: int = None, range_max: int = None, show: bool = False, custom_bins=None):
    from statisticFunctions import normal_dist_expected, uniform_dist_expected
    from matplotlib import pyplot as plt
    from matplotlib.text import Text
    from matplotlib.ticker import AutoMinorLocator
    import numpy as np

    if range_min is None:
        range_min = min(observed)
    if range_max is None:
        range_max = max(observed)

    # plot:
    fig, ax = plt.subplots(3, 1, figsize=(15, 20))

    bin_bounds = determine_bin_bounds(observed) if custom_bins is None else custom_bins

    _, bins, patches = ax[0].hist(observed, bins=bin_bounds, range=(range_min, range_max), ec='black')

    normal_expected = normal_dist_expected(observed, bins)  # Get the expected normal distribution
    uniform_expected = uniform_dist_expected(observed, bins)  # Get the expected normal distribution

    _, __, __ = ax[1].hist(normal_expected, bins=bins,  range=(range_min, range_max), ec='black')
    _, __, __ = ax[2].hist(uniform_expected, bins=bins,  range=(range_min, range_max), ec='black')

    if len(bins) == 1:
        plt.close()
        return None

    x_ticks = [(bins[idx + 1] + value) / 2 for idx, value in enumerate(bins[:-1])]
    x_ticks_labels = ["[{:.2f}-{:.2f})".format(value, bins[idx+1]) for idx, value in enumerate(bins[:-1])]
    x_ticks_labels[-1] = x_ticks_labels[-1][:-1] + "]"

    plt.setp(ax, xticks=x_ticks, xticklabels=x_ticks_labels)
    ax[0].tick_params(axis='x', which='minor', length=0)
    ax[1].tick_params(axis='x', which='minor', length=0)
    ax[2].tick_params(axis='x', which='minor', length=0)

    subplot_fontdict = {'fontsize': 19}
    ax[0].set_title('Observed', fontdict=subplot_fontdict)
    ax[1].set_title('Expected Normal Distribution', fontdict=subplot_fontdict)
    ax[2].set_title('Expected Uniform Distribution', fontdict=subplot_fontdict)

    filename = title.replace("\n", " ")
    plt_title = title.replace("\n", ", ")
    fig.suptitle(plt_title, size=26, y=0.95)
    plt.savefig(path + f"{filename}.png", bbox_inches="tight")

    plt.cla()
    if show:
        plt.show()
    plt.close()

    return [bins, observed, normal_expected, uniform_expected, plt_title]


def save_results_to_doc(runs: list, heading: str, path: str):
    from docx import Document
    from docx.shared import Inches

    document = Document()

    document.add_heading(heading, 0)

    for run in runs:
        p = document.add_paragraph()
        p.add_run(run["parameters"]).bold = True
        document.add_paragraph(run["chi_normal"], style='List Bullet')
        document.add_paragraph(run["chi_uniform"], style='List Bullet')

        document.add_picture(path + run["parameters"] + ".png", width=Inches(4))

        document.add_page_break()

    document.save(path + heading + ".docx")
